"""Resume CRUD（契约 §4.2）：POST/GET/PUT/DELETE + 列表 + 条目编辑锁定（§5.5）+ 重装配渲染（§6）+ 导出（§7 E8）。"""
import json
from typing import Optional

from fastapi import APIRouter, Depends, Request
from fastapi.responses import Response
from pydantic import BaseModel, Field

from ..core.errors import AppError, E_EXPORT, E_PARAM
from ..core.validation import check_resume
from ..engine.assembly import Assembler
from ..schemas import CamelModel, Resume

router = APIRouter(prefix="/api/resume", tags=["resume"])

# 可编辑条目块（§5.5）：summary 句 / internship duty / project item
EDITABLE_BLOCKS = ("summary", "internship", "project")


class ResumeIdResp(BaseModel):
    resume_id: str


class DeletedResp(BaseModel):
    deleted: bool


class ResumeListItem(BaseModel):
    id: str
    name: str = ""
    direction: str = ""
    updated_at: Optional[str] = None
    file: str = ""     # 本地存储位置（相对路径，供 UI 展示）


class ItemEditBody(CamelModel):
    """编辑锁定（§5.5）：summary 用 index；internship/project 用 index + sub_index 定位叶子。"""
    block: str = Field(min_length=1)                       # summary/internship/project
    index: int = Field(ge=0)                               # 父级下标（summary 为句子下标）
    sub_index: Optional[int] = Field(default=None, ge=0)   # 叶子下标（实习职责/项目要点）
    text: str = Field(min_length=1, max_length=500)


class ItemUnlockBody(CamelModel):
    block: str = Field(min_length=1)
    index: int = Field(ge=0)
    sub_index: Optional[int] = Field(default=None, ge=0)


class RenderBody(BaseModel):
    density: Optional[str] = Field(default=None, pattern="^(compact|normal|loose)$")


@router.post("", response_model=dict)
def create_resume(body: Resume, request: Request):
    """新建简历：无 id 则生成；集中校验后落库。"""
    storage = request.app.state.storage
    now = request.app.state.now
    if not body.id:
        body.id = storage.new_resume_id()
    body.created_at = body.created_at or now()
    body.updated_at = now()
    check_resume(body, request.app.state.config.limits)
    storage.save_resume(body.model_dump(mode="json", by_alias=True, exclude_none=False))
    return {"code": 0, "message": "ok", "data": {"resumeId": body.id}}


@router.get("", response_model=dict)
def list_resumes(request: Request):
    """简历列表（轻量）：前端工作台使用。"""
    storage = request.app.state.storage
    items = []
    for rid in storage.list_resumes():
        data = storage.load_resume(rid)
        name = (data.get("basicInfo") or {}).get("name", "")
        direction = data.get("direction") or ""
        items.append(ResumeListItem(id=rid, name=name, direction=direction,
                                    updated_at=data.get("updatedAt"),
                                    file=f"data/resumes/{rid}.json").model_dump())
    items.sort(key=lambda x: x["updated_at"] or "", reverse=True)
    return {"code": 0, "message": "ok", "data": {"items": items}}


@router.get("/{resume_id}", response_model=dict)
def get_resume(resume_id: str, request: Request):
    storage = request.app.state.storage
    data = storage.load_resume(resume_id)
    return {"code": 0, "message": "ok", "data": data}


@router.put("/{resume_id}", response_model=dict)
def update_resume(resume_id: str, body: Resume, request: Request):
    """整存更新：保留 id/created_at，刷新 updated_at。"""
    storage = request.app.state.storage
    now = request.app.state.now
    old = storage.load_resume(resume_id)
    body.id = resume_id
    body.created_at = body.created_at or old.get("created_at")
    body.updated_at = now()
    check_resume(body, request.app.state.config.limits)
    storage.save_resume(body.model_dump(mode="json", by_alias=True, exclude_none=False))
    return {"code": 0, "message": "ok", "data": {"updatedAt": body.updated_at}}


@router.delete("/{resume_id}", response_model=dict)
def delete_resume(resume_id: str, request: Request):
    storage = request.app.state.storage
    storage.load_resume(resume_id)  # 不存在 → 40008
    deleted = storage.delete_resume(resume_id)
    return {"code": 0, "message": "ok", "data": DeletedResp(deleted=deleted).model_dump()}


# ---------------------------------------------------------------- 编辑锁定（§5.5）与重装配（§6）

def _leaf(resume: dict, block: str, index: int, sub_index: Optional[int]) -> dict:
    """定位可编辑叶子条目：summary 句 / 实习职责 / 项目要点。非法定位 → 40001。"""
    if block == "summary":
        if sub_index is not None:
            raise AppError(E_PARAM, "summary 是单层列表，无需 subIndex", {"block": block})
        items = resume.get("summary") or []
        idx = index
    elif block in ("internship", "project"):
        parents = resume.get(block) or []
        if not (0 <= index < len(parents)):
            raise AppError(E_PARAM, f"{block} 下标越界", {"block": block, "index": index})
        key = "duties" if block == "internship" else "items"
        items = parents[index].get(key) or []
        if sub_index is None:
            raise AppError(E_PARAM, f"{block} 需要 subIndex 定位叶子", {"block": block})
        idx = sub_index
    else:
        raise AppError(E_PARAM, "不可编辑板块", {"block": block})
    if not (0 <= idx < len(items)):
        raise AppError(E_PARAM, "条目下标越界",
                       {"block": block, "index": index, "subIndex": sub_index})
    return items[idx]


def _rendered(resume: dict, app) -> dict:
    """重装配（§6）：density 取 resume 现值，返回 {resume, html, config} 供预览。"""
    assembler = Assembler(app.state.config.paths.templates_dir, app.state.storage)
    generation = resume.get("generation") or {}
    html, config = assembler.render(
        resume, {},
        density=str(resume.get("density") or "normal"),
        watermark_mode=str(generation.get("watermarkMode") or "practice"),
    )
    return {"code": 0, "message": "ok",
            "data": {"resume": resume, "html": html, "config": config}}


@router.put("/{resume_id}/item", response_model=dict)
def edit_item(resume_id: str, body: ItemEditBody, request: Request):
    """单条编辑（§5.5）：改文本 + edited:true + criticality 强制 critical → 不可被自动重写。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    leaf = _leaf(resume, body.block, body.index, body.sub_index)
    text = body.text.strip()
    if not text:
        raise AppError(E_PARAM, "文本不能为空", {"block": body.block})
    leaf["text"] = text[: (300 if body.block == "summary" else 500)]
    leaf["edited"] = True
    leaf["criticality"] = "critical"
    resume["updatedAt"] = request.app.state.now()
    storage.save_resume(resume)
    return _rendered(resume, request.app)


@router.post("/{resume_id}/item/unlock", response_model=dict)
def unlock_item(resume_id: str, body: ItemUnlockBody, request: Request):
    """解锁（§5.5）：edited:false → 下次自动生成可重写该条目。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    leaf = _leaf(resume, body.block, body.index, body.sub_index)
    leaf["edited"] = False
    resume["updatedAt"] = request.app.state.now()
    storage.save_resume(resume)
    return _rendered(resume, request.app)


@router.post("/{resume_id}/render", response_model=dict)
def render_resume(resume_id: str, body: RenderBody, request: Request):
    """重装配渲染（§6）：支持 density 手动调整后重新出图。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    if body.density:
        resume["density"] = body.density
        resume["updatedAt"] = request.app.state.now()
        storage.save_resume(resume)
    return _rendered(resume, request.app)


# ---------------------------------------------------------------- 导出（§7 E8：PDF / DOCX / JSON）


def _resume_to_docx(resume: dict) -> bytes:
    """最小可用 DOCX：结构化区块 + 要点列表（python-docx）。"""
    from io import BytesIO

    import docx

    document = docx.Document()
    info = resume.get("basicInfo") or {}
    document.add_heading(info.get("name") or "简历", level=0)
    contact = " | ".join(str(x) for x in
                         [info.get("phone"), info.get("email"), info.get("base"),
                          info.get("website")] if x)
    if contact:
        document.add_paragraph(contact)

    def section(title: str, lines: list) -> None:
        lines = [str(x).strip() for x in lines if str(x).strip()]
        if not lines:
            return
        document.add_heading(title, level=1)
        for ln in lines:
            document.add_paragraph(ln, style="List Bullet")

    section("自我评价", [s.get("text") for s in (resume.get("summary") or [])])
    section("教育经历", [
        f"{e.get('school')} · {e.get('major')}（{e.get('degree')}）"
        f"{e.get('startMonth')} - {e.get('endMonth')}" for e in (resume.get("education") or [])])
    for it in (resume.get("internship") or []):
        section(f"实习经历：{it.get('company')} · {it.get('position')}",
                [d.get("text") for d in (it.get("duties") or [])])
    for p in (resume.get("project") or []):
        tech = "、".join(p.get("techStack") or [])
        section(f"项目经验：{p.get('name')}（{tech}）",
                [x.get("text") for x in (p.get("items") or [])])
    section("技能特长", [s.get("name") for s in (resume.get("skill") or [])])
    section("证书荣誉", [h.get("name") for h in (resume.get("honor") or [])])

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def _resume_to_markdown(resume: dict) -> str:
    """导出 Markdown：结构化区块 + 无序列表（标准 GFM，兼容常见编辑器 / 知识库 / GitHub）。"""
    out = []
    info = resume.get("basicInfo") or {}
    out.append("# " + (info.get("name") or "简历"))
    contact = " | ".join(str(x) for x in
                         [info.get("phone"), info.get("email"), info.get("base"),
                          info.get("website")] if x)
    if contact:
        out.append("")
        out.append(contact)
    out.append("")

    def section(title: str, lines: list) -> None:
        lines = [str(x).strip() for x in lines if str(x).strip()]
        if not lines:
            return
        out.append("## " + title)
        out.extend("- " + ln for ln in lines)
        out.append("")

    section("自我评价", [s.get("text") for s in (resume.get("summary") or [])])
    section("教育经历", [
        f"{e.get('school')} · {e.get('major')}（{e.get('degree')}）"
        f"{e.get('startMonth')} - {e.get('endMonth')}" for e in (resume.get("education") or [])])
    for it in (resume.get("internship") or []):
        section(f"实习经历：{it.get('company')} · {it.get('position')}",
                [d.get("text") for d in (it.get("duties") or [])])
    for p in (resume.get("project") or []):
        tech = "、".join(p.get("techStack") or [])
        section(f"项目经验：{p.get('name')}（{tech}）",
                [x.get("text") for x in (p.get("items") or [])])
    section("技能特长", [s.get("name") for s in (resume.get("skill") or [])])
    section("证书荣誉", [h.get("name") for h in (resume.get("honor") or [])])
    return "\n".join(out).strip() + "\n"


def _resume_to_html(resume: dict) -> str:
    """导出标准 HTML5 文档：语义化区块 + 内嵌打印友好样式，可独立打开 / 浏览器打印。"""
    import html as _html

    esc = lambda x: _html.escape(str(x if x is not None else ""))
    info = resume.get("basicInfo") or {}
    body = ['<div class="resume">']
    body.append("<h1>" + esc(info.get("name") or "简历") + "</h1>")
    contact = " | ".join(str(x) for x in
                         [info.get("phone"), info.get("email"), info.get("base"),
                          info.get("website")] if x)
    if contact:
        body.append('<p class="contact">' + esc(contact) + "</p>")

    def section(title: str, lines: list) -> None:
        lines = [str(x).strip() for x in lines if str(x).strip()]
        if not lines:
            return
        body.append("<section><h2>" + esc(title) + "</h2><ul>")
        body.extend("<li>" + esc(ln) + "</li>" for ln in lines)
        body.append("</ul></section>")

    section("自我评价", [s.get("text") for s in (resume.get("summary") or [])])
    section("教育经历", [
        f"{e.get('school')} · {e.get('major')}（{e.get('degree')}）"
        f"{e.get('startMonth')} - {e.get('endMonth')}" for e in (resume.get("education") or [])])
    for it in (resume.get("internship") or []):
        section(f"实习经历：{it.get('company')} · {it.get('position')}",
                [d.get("text") for d in (it.get("duties") or [])])
    for p in (resume.get("project") or []):
        tech = "、".join(p.get("techStack") or [])
        section(f"项目经验：{p.get('name')}（{tech}）",
                [x.get("text") for x in (p.get("items") or [])])
    section("技能特长", [s.get("name") for s in (resume.get("skill") or [])])
    section("证书荣誉", [h.get("name") for h in (resume.get("honor") or [])])
    body.append("</div>")
    css = (
        "body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,"
        "'PingFang SC','Microsoft YaHei',sans-serif;max-width:820px;margin:0 auto;"
        "padding:32px 24px;color:#1f2937;line-height:1.7}"
        "h1{font-size:26px;margin:0 0 4px}.contact{color:#6b7280;font-size:14px;margin:0}"
        "section{margin-top:22px}h2{font-size:17px;border-bottom:1px solid #e5e7eb;padding-bottom:6px;margin:0}"
        "ul{margin:10px 0 0;padding-left:20px}li{margin-bottom:5px}"
        "@media print{body{padding:0}}"
    )
    title = esc(info.get("name") or "简历")
    return (
        "<!DOCTYPE html>\n<html lang=\"zh-CN\">\n<head>\n<meta charset=\"utf-8\">\n"
        "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">\n"
        "<title>" + title + " - 简历</title>\n<style>" + css + "</style>\n</head>\n<body>\n"
        + "\n".join(body) + "\n</body>\n</html>\n"
    )


@router.get("/{resume_id}/export", response_model=dict)
def export_resume(resume_id: str, request: Request, format: str = "json"):
    """导出：format=json（结构化数据）/ docx（Word）/ md|markdown / html / pdf（由前端打印生成）。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    fmt = format.lower()
    if fmt == "json":
        data = json.dumps(resume, ensure_ascii=False, indent=2).encode("utf-8")
        return Response(
            content=data, media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.json"'})
    if fmt == "docx":
        try:
            content = _resume_to_docx(resume)
        except ImportError:
            raise AppError(E_EXPORT, "DOCX 导出需要安装 python-docx（pip install python-docx）")
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.docx"'})
    if fmt in ("md", "markdown"):
        content = _resume_to_markdown(resume).encode("utf-8")
        return Response(
            content=content, media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.md"'})
    if fmt == "html":
        content = _resume_to_html(resume).encode("utf-8")
        return Response(
            content=content, media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.html"'})
    raise AppError(E_PARAM, "不支持的导出格式", {"format": fmt})
